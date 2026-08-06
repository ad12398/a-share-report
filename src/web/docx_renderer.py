"""Word 文档渲染器 —— 将 HTML 报告转换为 .docx（python-docx）"""

import logging
import re
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

logger = logging.getLogger("a-share-report")

PROJECT_ROOT = Path(__file__).parent.parent.parent
OUTPUT_DIR = PROJECT_ROOT / "reports"

# 颜色定义
COLOR_HEADING = RGBColor(0x1F, 0x77, 0xB4)    # 深蓝
COLOR_BODY = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_ACCENT = RGBColor(0xD2, 0x99, 0x22)       # 橙色（重点文字）
COLOR_MUTED = RGBColor(0x8B, 0x94, 0x9E)        # 灰色


class _HTMLToDocxParser(HTMLParser):
    """将 HTML 解析为 python-docx Document"""

    def __init__(self, doc: Document):
        super().__init__()
        self.doc = doc
        self.para = None        # 当前段落
        self.list_level = 0     # 当前列表嵌套层数（0=不在列表中）
        self.current_tag = None
        self.bold_stack: list[bool] = []  # 粗体标签栈

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]):
        tag = tag.lower()
        if tag == "h3":
            self.para = self.doc.add_heading(level=2)
            self.current_tag = "h3"
        elif tag == "h4":
            self.para = self.doc.add_heading(level=3)
            self.current_tag = "h4"
        elif tag == "p":
            self.para = self.doc.add_paragraph()
            self.para.paragraph_format.space_after = Pt(6)
            self.current_tag = "p"
        elif tag == "li":
            # li 在 ul/ol 内部
            self.para = self.doc.add_paragraph(style="List Bullet")
            self.para.paragraph_format.space_after = Pt(2)
            self.current_tag = "li"
        elif tag == "ul" or tag == "ol":
            pass  # python-docx 自动处理列表，无需额外操作
        elif tag == "br":
            if self.para:
                self.para.add_run().add_break()
        elif tag == "strong":
            self.bold_stack.append(True)

    def handle_endtag(self, tag: str):
        tag = tag.lower()
        if tag in ("h3", "h4", "p", "li"):
            self.para = None
            self.current_tag = None
        elif tag == "strong":
            if self.bold_stack:
                self.bold_stack.pop()

    def handle_data(self, data: str):
        text = data.strip()
        if not text:
            return
        if self.para is None:
            # 浮动文本（不在任何段落标签内），创建一个新段落
            self.para = self.doc.add_paragraph()
            self.para.paragraph_format.space_after = Pt(6)

        is_bold = any(self.bold_stack)
        run = self.para.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_BODY
        if is_bold:
            run.bold = True
            run.font.color.rgb = COLOR_ACCENT

        # 设置中文字体
        run.font.name = "微软雅黑"
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def _preprocess_html(html_text: str) -> str:
    """预处理 HTML：清理 <br> 标签、修复常见格式问题"""
    # 将 <br> 转为段落分隔（如果两个 <br> 之间无内容，分段落）
    text = html_text.replace("<br>", "\n<br>\n")
    # 确保标签是完整 HTML
    if "<html" not in text.lower():
        text = f"<div>{text}</div>"
    return text


def render_docx(
    slot: str,
    slot_label: str,
    date_str: str,
    report_html: str,
    index_data: dict[str, Any],
    movers: dict[str, Any],
    north_flow: dict[str, Any],
    fund_flow: dict[str, Any],
    macro: dict[str, Any],
) -> Document:
    """将报告渲染为 Word Document 对象"""

    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # ── 封面标题 ──
    title = doc.add_heading(f"A 股量化报告 — {slot_label}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"{date_str}  |  模型: DeepSeek  |  量化交易视角")
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_MUTED

    # ── KPI 摘要表 ──
    doc.add_heading("核心指标速览", level=2)

    # 指数行情
    if index_data:
        idx_table = doc.add_table(rows=1, cols=4, style="Light Shading Accent 1")
        hdr = idx_table.rows[0].cells
        hdr[0].text = "指数"; hdr[1].text = "点位"; hdr[2].text = "涨跌幅"; hdr[3].text = ""
        for code, info in index_data.items():
            if code == "_validation":
                continue
            row = idx_table.add_row().cells
            row[0].text = info.get("name", code)
            row[1].text = f"{info.get('price', 0):.2f}"
            pct = info.get("change_pct", 0)
            row[2].text = f"{'+' if pct >= 0 else ''}{pct:.2f}%"
            row[3].text = ""
        doc.add_paragraph()  # 间距

    # 外资监测（活跃度+南向+外资占比）
    if north_flow and north_flow.get("turnover_total"):
        # 标题行
        p = doc.add_paragraph()
        p.add_run("外资监测").bold = True

        # ① 北向活跃度
        turnover = north_flow.get("turnover_total", 0)
        p = doc.add_paragraph()
        level = "高" if turnover > 300 else "中" if turnover > 150 else "低"
        p.add_run(f"  北向活跃度：成交 {turnover:.0f} 亿（{level}活跃度）")

        # ② 外资占比
        participation = north_flow.get("participation_pct", 0)
        if participation > 0:
            p = doc.add_paragraph()
            p.add_run(f"  外资占比：{participation:.1f}%（北向成交 / 两市成交）")

        # ③ 沪/深偏好
        sh_ratio = north_flow.get("sh_ratio", 0)
        if sh_ratio > 0:
            preference = "偏价值防御" if sh_ratio > 55 else "偏成长进攻" if sh_ratio < 45 else "均衡"
            p = doc.add_paragraph()
            p.add_run(f"  沪/深偏好：沪市 {sh_ratio:.0f}%（{preference}）")

        # ④ 南向资金
        south = north_flow.get("south_flow", {}) or {}
        south_net = south.get("south_net", 0) or 0
        if south_net:
            south_dir = "净买入" if south_net > 0 else "净卖出"
            p = doc.add_paragraph()
            p.add_run(f"  南向资金：{south_dir} {abs(south_net):.1f} 亿（港股通，A股情绪反向指标）")

        # 数据限制
        run = p.add_run(f"\n  注：北向净买入自2024年证监会新规后不再公开发布，方向由活跃度+南向+外部联动综合推断")
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_MUTED

    # 资金流（替代两融）
    if fund_flow and fund_flow.get("sample_count"):
        p = doc.add_paragraph()
        p.add_run(f"大资金方向（替代两融）：").bold = True
        p.add_run(
            f"{fund_flow['sample_count']} 股聚合净流入 {fund_flow['total_net_amount']} 亿  |  "
            f"大单净 {fund_flow['big_order_net']} 亿  |  "
            f"涨家 {fund_flow['inflow_count']} / 跌家 {fund_flow['outflow_count']}"
        )
        run = p.add_run(f"\n注：{fund_flow.get('_note', '')}")
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_MUTED

    # 宏观数据
    if macro and macro.get("cpi"):
        p = doc.add_paragraph()
        cpi = macro["cpi"]
        ppi = macro.get("ppi", {})
        pmi = macro.get("pmi", {})
        m2 = macro.get("m2", {})
        lpr = macro.get("lpr", {})
        p.add_run("宏观数据：").bold = True
        parts = []
        if cpi.get("yoy"): parts.append(f"CPI {cpi.get('yoy', '-')}%")
        if ppi.get("yoy"): parts.append(f"PPI {ppi.get('yoy', '-')}%")
        if pmi.get("manufacturing"): parts.append(f"PMI {pmi.get('manufacturing', '-')}")
        if m2.get("yoy"): parts.append(f"M2 {m2.get('yoy', '-')}%")
        if lpr.get("lpr_1y"): parts.append(f"LPR 1Y {lpr.get('lpr_1y', '-')}%")
        p.add_run("  |  ".join(parts))

    doc.add_paragraph()  # 间距

    # ── AI 分析正文 ──
    doc.add_heading("量化分析", level=2)

    html_text = _preprocess_html(report_html)
    parser = _HTMLToDocxParser(doc)
    try:
        parser.feed(html_text)
    except Exception as e:
        logger.warning(f"HTML 解析部分失败: {e}")
        # 降级：纯文本
        clean = re.sub(r"<[^>]+>", "", report_html)
        doc.add_paragraph(clean)

    # ── 免责声明 ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "⚠️ 免责声明：本报告由 AI 自动生成，仅供参考，不构成任何投资建议。"
        "投资有风险，入市需谨慎。报告中的数据来源于公开接口，不对其准确性做任何保证。"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_MUTED

    return doc


def save_docx(doc: Document, slot: str, date_str: str) -> str:
    """保存 Word 文档到 reports 目录，返回相对路径

    文件名格式: {date} {slot_hour}时{slot_min}分 A股量化报告.docx
    例如: 2026-08-03 15时00分 A股量化报告.docx
    """
    report_dir = OUTPUT_DIR / date_str
    report_dir.mkdir(parents=True, exist_ok=True)

    slot_hour = slot[:2] if len(slot) >= 2 else slot
    slot_min = slot[2:] if len(slot) >= 4 else "00"
    filename = f"{date_str} {slot_hour}时{slot_min}分 A股量化报告.docx"
    filepath = report_dir / filename
    doc.save(str(filepath))
    logger.info(f"Word 文档已保存: {filepath}")
    return str(filepath.relative_to(PROJECT_ROOT))
